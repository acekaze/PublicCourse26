from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT_DIR = Path(r"C:\코딩\공개과정\output\doc")
OUT_PATH = OUT_DIR / "주력강의재설계과정_표지_최종본.docx"
LOGO_PATH = Path(r"C:\Users\aceka\Downloads\전환설계연구소_2 (1).png")

NAVY = RGBColor(23, 55, 117)
GRAY = RGBColor(90, 104, 128)
BLACK = RGBColor(20, 20, 20)
FONT_BODY = "Pretendard Variable"
FONT_HEAD = "Pretendard Variable Black"


def apply_font(run, font_name: str) -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:ascii"), font_name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font_name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run._element.rPr.rFonts.set(qn("w:cs"), font_name)


def styled_run(paragraph, text: str, *, size=10.5, bold=False, color=None, font_name=None):
    run = paragraph.add_run(text)
    apply_font(run, font_name or FONT_BODY)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    return run


def build_document() -> Document:
    document = Document()
    section = document.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    normal = document.styles["Normal"]
    normal.font.name = FONT_BODY
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT_BODY)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_BODY)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
    normal._element.rPr.rFonts.set(qn("w:cs"), FONT_BODY)
    normal.font.size = Pt(10.5)

    top = document.add_paragraph()
    top.alignment = WD_ALIGN_PARAGRAPH.CENTER
    top.paragraph_format.space_before = Pt(26)
    styled_run(top, "전환설계연구소", size=12, color=GRAY, font_name=FONT_BODY)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(90)
    styled_run(title, "주력 강의 재설계 과정", size=28, bold=True, color=NAVY, font_name=FONT_HEAD)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_before = Pt(8)
    styled_run(subtitle, "참가자 워크북", size=15, bold=True, color=BLACK, font_name=FONT_HEAD)

    if LOGO_PATH.exists():
        logo = document.add_paragraph()
        logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        logo.paragraph_format.space_before = Pt(110)
        run = logo.add_run()
        run.add_picture(str(LOGO_PATH), width=Cm(16.8))

    footer = document.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.paragraph_format.space_before = Pt(120)
    styled_run(footer, "자기 강의를 실제로 손보고 개선안을 만드는 2일 과정", size=10.5, color=GRAY, font_name=FONT_BODY)

    return document


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    document = build_document()
    document.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    main()
