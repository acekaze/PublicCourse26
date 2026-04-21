from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT_DIR = Path(r"C:\코딩\공개과정\output\doc")
OUT_PATH = OUT_DIR / "주력강의재설계과정_2일커리큘럼_시간표.docx"


def set_doc_defaults(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    normal = document.styles["Normal"]
    normal.font.name = "Malgun Gothic"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    normal.font.size = Pt(10.5)

    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        style = document.styles[style_name]
        style.font.name = "Malgun Gothic"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_title(document: Document) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("주력 강의 재설계 과정\n2일 커리큘럼 · 시간표")
    run.bold = True
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(23, 55, 117)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    srun = subtitle.add_run("2026년 5월 1일(금) - 5월 2일(토) · 문래오층")
    srun.font.name = "Malgun Gothic"
    srun._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    srun.font.size = Pt(10.5)
    srun.font.color.rgb = RGBColor(90, 104, 128)


def add_bullets(document: Document, title: str, items: list[str]) -> None:
    heading = document.add_paragraph()
    hr = heading.add_run(title)
    hr.bold = True
    hr.font.name = "Malgun Gothic"
    hr._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    hr.font.size = Pt(12)
    hr.font.color.rgb = RGBColor(23, 55, 117)

    for item in items:
        p = document.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        r.font.name = "Malgun Gothic"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")


def add_schedule_table(document: Document, day_title: str, focus: str, rows: list[list[str]]) -> None:
    title = document.add_paragraph()
    title_run = title.add_run(day_title)
    title_run.bold = True
    title_run.font.name = "Malgun Gothic"
    title_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    title_run.font.size = Pt(14)
    title_run.font.color.rgb = RGBColor(23, 55, 117)

    focus_p = document.add_paragraph()
    focus_run = focus_p.add_run(focus)
    focus_run.font.name = "Malgun Gothic"
    focus_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    focus_run.font.size = Pt(10.5)

    table = document.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    headers = ["시간", "세션", "주요 내용", "산출물"]
    widths = [Cm(2.4), Cm(4.2), Cm(8.1), Cm(3.8)]
    hdr_cells = table.rows[0].cells
    for idx, text in enumerate(headers):
        hdr_cells[idx].width = widths[idx]
        hdr_cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(hdr_cells[idx], "DCE6F8")
        p = hdr_cells[idx].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = True
        run.font.name = "Malgun Gothic"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        run.font.size = Pt(10)

    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].width = widths[idx]
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[idx].paragraphs[0]
            if idx == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(value)
            run.font.name = "Malgun Gothic"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
            run.font.size = Pt(9.5)


def build_document() -> Document:
    document = Document()
    set_doc_defaults(document)
    add_title(document)

    document.add_paragraph("")

    add_bullets(
        document,
        "과정 운영 기준",
        [
            "1일차는 누구를 위한 강의인지, 무엇을 통해 어떤 변화를 만드는 강의인지를 명확히 하고 전달력 플로우 구조로 설계하는 날입니다.",
            "2일차는 학습자 촉진과 전달기술을 점검하고 종합실습을 통해 주력 강의 개선안을 정리하는 날입니다.",
            "설명보다 작성, 공유, 피드백, 수정의 비중이 높습니다.",
        ],
    )

    document.add_paragraph("")

    day1_rows = [
        ["10:00-11:00", "오프닝 · 아이스브레이킹 · 연결 경험", "강의에 대한 연결 경험을 만들고 현재 강의 상태와 환경설정 기준을 함께 점검합니다.", "현재 강의 상태 메모"],
        ["11:00-12:00", "WHO 점검", "누구를 위한 강의인지, 학습자는 어떤 상태와 과제를 가지고 들어오는지 정리합니다.", "대상자 정의서"],
        ["12:00-13:00", "핵심 콘텐츠 명확화", "질문을 통해 강사의 핵심 콘텐츠와 변화 포인트를 압축합니다.", "핵심 콘텐츠 메모"],
        ["14:00-15:00", "강의 정의 문장 완성", "나는 누구에게 어떠한 내용을 통해 어떠한 변화를 줄 수 있는 강의를 한다 문장을 완성합니다.", "강의 정의 문장"],
        ["15:00-16:00", "전달력 플로우 이해 및 적용", "전달력 플로우를 이해하고 자기 강의에 적용할 기준을 잡습니다.", "전달력 플로우 설계 프레임"],
        ["16:00-17:00", "전달력 플로우 기반 1차 설계", "완성한 핵심 문장을 기준으로 주력 강의를 전달력 플로우 구조로 재배치합니다.", "1차 강의 구조안"],
        ["17:00-18:00", "피드백 및 정리", "구조안과 메시지의 선명도를 점검하고 2일차 보완 지점을 정리합니다.", "피드백 메모"],
    ]
    add_schedule_table(
        document,
        "Day 1",
        "핵심 목표: 대상, 핵심 콘텐츠, 변화 목표를 명확히 하고 전달력 플로우 구조로 1차 설계를 완성합니다.",
        day1_rows,
    )

    document.add_paragraph("")

    day2_rows = [
        ["10:00-11:00", "Day 1 리뷰 및 학습자 촉진 1", "전일 결과를 복기하고 질문, 사례, 활동이 학습 목표와 어떻게 연결되는지 점검합니다.", "촉진 설계 메모"],
        ["11:00-12:00", "학습자 촉진 2", "참여를 만드는 질문 설계, 활동 배치, 피드백 운영 방식을 구체화합니다.", "질문 · 활동 설계안"],
        ["12:00-13:00", "전달기술 1", "설명, 전환, 강조, 마무리의 전달 방식을 점검합니다.", "전달기술 점검표"],
        ["14:00-15:00", "전달기술 2", "강의 흐름 안에서 전달력이 떨어지는 구간을 확인하고 수정합니다.", "전달 수정 메모"],
        ["15:00-16:00", "종합실습 · 미니 시연 1", "핵심 구간을 실제로 시연하고 메시지, 구조, 촉진, 전달 관점에서 피드백을 받습니다.", "미니 시연 피드백 기록"],
        ["16:00-17:00", "종합실습 · 미니 시연 2", "피드백을 반영해 강의 운영 방식을 더 구체적으로 다듬습니다.", "수정 반영 메모"],
        ["17:00-18:00", "개선안 정리 및 마무리", "주력 강의 개선안을 정리하고 현장 적용 포인트를 확인합니다.", "주력 강의 개선안"],
    ]
    add_schedule_table(
        document,
        "Day 2",
        "핵심 목표: 학습자 촉진과 전달기술을 점검하고 종합실습을 통해 주력 강의 개선안을 정리합니다.",
        day2_rows,
    )

    document.add_paragraph("")
    add_bullets(
        document,
        "참가자 안내용 핵심 문구",
        [
            "1일차는 자기 강의의 대상, 핵심 콘텐츠, 변화 목표를 명확히 하고 전달력 플로우 구조로 설계하는 날입니다.",
            "2일차는 학습자 촉진과 전달기술을 점검하고 실제 시연과 피드백을 통해 주력 강의 개선안을 정리하는 날입니다.",
            "이 과정은 많이 배우고 가는 과정이 아니라 자기 강의를 실제로 손보고 가는 과정입니다.",
        ],
    )

    return document


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    document = build_document()
    document.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    main()
