from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT_DIR = Path(r"C:\코딩\공개과정\output\doc")
OUT_PATH = OUT_DIR / "주력강의재설계과정_워크북_1차시안.docx"
ILLUS_DIR = Path(r"C:\코딩\공개과정\output\doc\illustrations")
IMG_FLOW = ILLUS_DIR / "전달력플로우_삽화.png"
IMG_ENGAGEMENT = ILLUS_DIR / "참여도흐름_삽화.png"
IMG_ENERGY = ILLUS_DIR / "에너지흐름과액티비티설계_삽화_v2.png"
IMG_PAIRED = ILLUS_DIR / "PairedShare_삽화.png"
IMG_MOVEMENT = ILLUS_DIR / "동선_삽화_v2.png"
IMG_ANCHOR = ILLUS_DIR / "로케이션앵커링_삽화_v2.png"
IMG_NONVERBAL = ILLUS_DIR / "비언어전달_삽화_v3.png"

NAVY = RGBColor(23, 55, 117)
GRAY = RGBColor(90, 104, 128)
LIGHT_BLUE = "EAF0FB"
PALE = "FBFCFE"
RULE = RGBColor(210, 220, 236)
FONT_BODY = "Pretendard Variable"
FONT_HEAD = "Pretendard Variable Black"


def apply_font(run, font_name: str) -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:ascii"), font_name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font_name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run._element.rPr.rFonts.set(qn("w:cs"), font_name)


def set_doc_defaults(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    normal = document.styles["Normal"]
    normal.font.name = FONT_BODY
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT_BODY)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_BODY)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
    normal._element.rPr.rFonts.set(qn("w:cs"), FONT_BODY)
    normal.font.size = Pt(10.5)

    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        style = document.styles[style_name]
        style.font.name = FONT_HEAD
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT_HEAD)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_HEAD)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_HEAD)
        style._element.rPr.rFonts.set(qn("w:cs"), FONT_HEAD)

    footer_p = section.footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    styled_run(footer_p, "주력 강의 재설계 과정 워크북", size=9, color=GRAY, font_name=FONT_BODY)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_borders(
    cell,
    *,
    top: bool = True,
    left: bool = True,
    bottom: bool = True,
    right: bool = True,
    color: str = "C8D4E8",
    size: str = "8",
) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    edge_map = {
        "top": top,
        "left": left,
        "bottom": bottom,
        "right": right,
    }
    for edge, enabled in edge_map.items():
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single" if enabled else "nil")
        element.set(qn("w:sz"), size if enabled else "0")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color if enabled else "FFFFFF")


def set_row_height(row, height_twips: int) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tr_height = OxmlElement("w:trHeight")
    tr_height.set(qn("w:val"), str(height_twips))
    tr_height.set(qn("w:hRule"), "atLeast")
    tr_pr.append(tr_height)


def styled_run(paragraph, text: str, size=10.5, bold=False, color=None, font_name: str | None = None):
    run = paragraph.add_run(text)
    apply_font(run, font_name or FONT_BODY)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    return run


def add_page_break(document: Document) -> None:
    document.add_page_break()


def add_rule(document: Document) -> None:
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run("_" * 60)
    run.font.color.rgb = RULE
    run.font.size = Pt(8)


def add_illustration_placeholder(document: Document, label: str, caption: str, height_cm: float = 7.0) -> None:
    table = document.add_table(rows=2, cols=1)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    top_row = table.rows[0]
    set_row_height(top_row, int(height_cm * 567))
    top_cell = top_row.cells[0]
    set_cell_shading(top_cell, "F5F8FE")
    p = top_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(80)
    styled_run(p, "삽화 삽입 영역\n", size=16, bold=True, color=GRAY, font_name=FONT_HEAD)
    styled_run(p, label, size=11, color=GRAY, font_name=FONT_BODY)

    bottom_cell = table.rows[1].cells[0]
    set_cell_shading(bottom_cell, PALE)
    cp = bottom_cell.paragraphs[0]
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    styled_run(cp, caption, size=9.5, color=GRAY, font_name=FONT_BODY)


def add_illustration_image(document: Document, image_path: Path, caption: str, width_cm: float = 15.8) -> None:
    if not image_path.exists():
        add_illustration_placeholder(document, "이미지 파일 누락", caption)
        return

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))

    cp = document.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    styled_run(cp, caption, size=9.5, color=GRAY)


def add_cover(document: Document) -> None:
    banner = document.add_table(rows=1, cols=1)
    banner.style = "Table Grid"
    banner.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = banner.rows[0].cells[0]
    set_row_height(banner.rows[0], 1900)
    set_cell_shading(cell, "E8F0FF")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(30)
    styled_run(p, "주력 강의 재설계 과정\n", size=24, bold=True, color=NAVY, font_name=FONT_HEAD)
    styled_run(p, "참가자 워크북", size=17, bold=True, color=NAVY, font_name=FONT_HEAD)

    info = document.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.paragraph_format.space_before = Pt(16)
    styled_run(info, "2026년 5월 1일(금) - 5월 2일(토)\n", size=11, color=GRAY, font_name=FONT_BODY)
    styled_run(info, "문래오층", size=11, color=GRAY, font_name=FONT_BODY)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_before = Pt(14)
    styled_run(subtitle, "자기 강의를 실제로 손보고 개선안을 만드는 2일 워크북", size=11, color=GRAY, font_name=FONT_BODY)

    card = document.add_table(rows=4, cols=2)
    card.alignment = WD_TABLE_ALIGNMENT.CENTER
    card.style = "Table Grid"
    labels = ["이름", "소속", "연락처", "이메일"]
    for i, label in enumerate(labels):
        left = card.rows[i].cells[0]
        right = card.rows[i].cells[1]
        set_cell_shading(left, LIGHT_BLUE)
        set_cell_shading(right, PALE)
        set_row_height(card.rows[i], 560)
        lp = left.paragraphs[0]
        lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        styled_run(lp, label, size=10, bold=True, color=NAVY, font_name=FONT_HEAD)
        styled_run(right.paragraphs[0], " ", size=10, font_name=FONT_BODY)


def add_page_header(document: Document, title: str, subtitle: str | None = None) -> None:
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    styled_run(p, title, size=18, bold=True, color=NAVY, font_name=FONT_HEAD)
    if subtitle:
        sp = document.add_paragraph()
        sp.paragraph_format.space_after = Pt(8)
        styled_run(sp, subtitle, size=10.5, color=GRAY, font_name=FONT_BODY)
    add_rule(document)


def add_section_divider(document: Document, label: str, title: str, body: str) -> None:
    box = document.add_table(rows=1, cols=1)
    box.style = "Table Grid"
    box.alignment = WD_TABLE_ALIGNMENT.CENTER
    row = box.rows[0]
    set_row_height(row, 2400)
    cell = row.cells[0]
    set_cell_shading(cell, "F2F6FD")
    p1 = cell.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.space_before = Pt(42)
    styled_run(p1, label + "\n", size=10, bold=True, color=GRAY, font_name=FONT_HEAD)
    styled_run(p1, title, size=20, bold=True, color=NAVY, font_name=FONT_HEAD)
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(10)
    styled_run(p2, body, size=11, color=GRAY, font_name=FONT_BODY)


def add_subsection_banner(document: Document, title: str, body: str) -> None:
    box = document.add_table(rows=1, cols=1)
    box.style = "Table Grid"
    box.alignment = WD_TABLE_ALIGNMENT.CENTER
    row = box.rows[0]
    set_row_height(row, 1100)
    cell = row.cells[0]
    set_cell_shading(cell, "F2F6FD")
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(12)
    styled_run(p, title + "\n", size=15, bold=True, color=NAVY, font_name=FONT_HEAD)
    styled_run(p, body, size=10, color=GRAY, font_name=FONT_BODY)


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        p = document.add_paragraph(style="List Bullet")
        styled_run(p, item, size=10.5, font_name=FONT_BODY)


def add_write_box(document: Document, label: str, height_cm: float = 3.0) -> None:
    label_p = document.add_paragraph()
    styled_run(label_p, label, size=10.5, bold=True, color=NAVY, font_name=FONT_HEAD)
    table = document.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    row = table.rows[0]
    set_row_height(row, int(height_cm * 567))
    cell = row.cells[0]
    cell.width = Cm(17.2)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    set_cell_shading(cell, PALE)
    set_cell_borders(cell)
    styled_run(cell.paragraphs[0], " ", size=10, font_name=FONT_BODY)


def add_blank_write_box(document: Document, height_cm: float = 3.0) -> None:
    table = document.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    row = table.rows[0]
    set_row_height(row, int(height_cm * 567))
    cell = row.cells[0]
    cell.width = Cm(17.2)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    set_cell_shading(cell, PALE)
    set_cell_borders(cell)
    styled_run(cell.paragraphs[0], " ", size=10, font_name=FONT_BODY)


def add_memo_pages(document: Document, count: int = 2) -> None:
    for i in range(count):
        add_page_break(document)
        add_page_header(document, "Memo", "과정 중 떠오른 생각, 적용 아이디어, 질문을 자유롭게 기록합니다.")
        table = document.add_table(rows=18, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        for row in table.rows:
            set_row_height(row, 480)
            cell = row.cells[0]
            cell.width = Cm(17.2)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.BOTTOM
            set_cell_shading(cell, "FFFFFF")
            set_cell_borders(
                cell,
                top=False,
                left=False,
                right=False,
                bottom=True,
                color="D5DDEA",
                size="6",
            )
            styled_run(cell.paragraphs[0], " ", size=10, font_name=FONT_BODY)


def add_two_column_boxes(document: Document, labels: list[str], height_cm: float = 3.2) -> None:
    table = document.add_table(rows=len(labels), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for i, label in enumerate(labels):
        row = table.rows[i]
        set_row_height(row, int(height_cm * 567))
        left, right = row.cells
        left.width = Cm(4.2)
        right.width = Cm(11.8)
        left.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        right.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        set_cell_shading(left, LIGHT_BLUE)
        set_cell_shading(right, PALE)
        set_cell_borders(left)
        set_cell_borders(right)
        lp = left.paragraphs[0]
        lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        styled_run(lp, label, size=10, bold=True, color=NAVY, font_name=FONT_HEAD)
        styled_run(right.paragraphs[0], " ", size=10, font_name=FONT_BODY)


def add_two_col_notes(document: Document, left_label: str, right_label: str, height_cm: float = 3.0) -> None:
    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    row = table.rows[0]
    set_row_height(row, int(height_cm * 567))
    labels = [left_label, right_label]
    for idx, cell in enumerate(row.cells):
        cell.width = Cm(8.45)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        set_cell_shading(cell, PALE)
        set_cell_borders(cell)
        p = cell.paragraphs[0]
        styled_run(p, labels[idx] + "\n", size=10, bold=True, color=NAVY, font_name=FONT_HEAD)
        styled_run(p, " ", size=10, font_name=FONT_BODY)


def add_check_table(document: Document, rows: list[str], cols: list[str]) -> None:
    table = document.add_table(rows=1 + len(rows), cols=1 + len(cols))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    hdr = table.rows[0].cells
    set_cell_shading(hdr[0], LIGHT_BLUE)
    set_cell_borders(hdr[0])
    styled_run(hdr[0].paragraphs[0], "점검 항목", size=9.5, bold=True, color=NAVY, font_name=FONT_HEAD)
    for i, col in enumerate(cols, start=1):
        set_cell_shading(hdr[i], LIGHT_BLUE)
        set_cell_borders(hdr[i])
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        styled_run(p, col, size=9.5, bold=True, color=NAVY, font_name=FONT_HEAD)

    for r_idx, row_label in enumerate(rows, start=1):
        row = table.rows[r_idx].cells
        set_cell_shading(row[0], "F6F8FC")
        set_cell_borders(row[0])
        styled_run(row[0].paragraphs[0], row_label, size=9.5, font_name=FONT_BODY)
        for c_idx in range(1, len(cols) + 1):
            set_cell_borders(row[c_idx])
            row[c_idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            styled_run(row[c_idx].paragraphs[0], "□", size=12, color=GRAY, font_name=FONT_BODY)


def build_document() -> Document:
    document = Document()
    set_doc_defaults(document)

    add_cover(document)
    add_page_break(document)

    add_page_header(document, "이 워크북을 쓰는 방법", "실습과 수정 중심으로 사용하는 참가자용 워크북입니다.")
    add_bullets(document, [
        "현재 운영 중인 주력 강의 1개를 기준으로 작성합니다.",
        "완벽하게 쓰기보다 수정 가능한 상태를 만드는 데 집중합니다.",
        "피드백은 들은 뒤 바로 기록합니다.",
        "과정 종료 시 강의 정의 문장, 전달력 플로우 구조안, 질문·활동 설계안, 주력 강의 개선안을 남깁니다.",
    ])
    add_write_box(document, "이번 과정에서 가장 먼저 점검하고 싶은 것", 3.2)
    add_write_box(document, "이번 과정에서 가장 얻고 싶은 것", 3.2)

    add_page_break(document)
    add_page_header(document, "학습 프레임과 기대사항", "학습 태도를 정렬하고, 이번 과정에서 얻고 싶은 것을 분명히 합니다.")
    add_bullets(document, [
        "새로운 것을 배우고 접할 때 불편하거나 혼란스러운 느낌이 드는 것은 자연스러운 과정입니다.",
        "학습은 실제 적용과 연결될 때 가장 빠르게 일어납니다.",
        "강사는 모든 현장 맥락의 전문가일 수 없으므로 참가자의 사례 공유가 중요합니다.",
        "참여도와 현장 적용 수준은 비례합니다.",
    ])
    add_write_box(document, "이번 과정에서 내가 가장 경계해야 할 태도", 2.8)
    add_write_box(document, "이번 과정에서 내가 가장 강화해야 할 태도", 2.8)
    add_write_box(document, "내 주변 사람들은 내가 이 과정을 통해 무엇을 가져오길 기대할까", 3.2)

    add_page_break(document)
    add_page_header(document, "오프닝 자기 점수", "강의 시작 전 현재 상태를 짧게 점검하고, 이후 변화를 비교할 기준을 만듭니다.")
    add_write_box(document, "나의 오늘 점수", 2.2)
    add_write_box(document, "그 이유", 3.2)

    add_page_break(document)
    add_page_header(document, "강사로서의 강점과 원하는 교수자 이미지", "지금의 강점과 앞으로 보이고 싶은 강사의 모습을 함께 정리합니다.")
    add_two_col_notes(document, "내가 가진 강점", "어떤 교수자로 보이길 원하는가", 4.0)
    add_write_box(document, "지금 내 강점이 가장 잘 드러나는 순간", 3.0)

    add_page_break(document)
    add_section_divider(
        document,
        "PART 1",
        "과정 재설계",
        "대상과 핵심을 명확히 하고, 강의 정의 문장과 전달력 플로우 기반 1차 구조안을 만듭니다.",
    )

    add_page_break(document)
    add_page_header(document, "현재 강의 상태와 연결 경험", "지금 가져온 강의의 상태와 도입부 연결 지점을 함께 점검합니다.")
    add_two_column_boxes(document, [
        "주력 강의명",
        "대상",
        "강의 시간",
        "운영 맥락",
    ], height_cm=2.0)
    add_write_box(document, "현재 보유 자료", 2.8)
    add_write_box(document, "지금 가장 먼저 수정하고 싶은 점", 2.8)
    add_write_box(document, "현재 도입부 점검 및 연결 경험 메모", 3.5)

    add_page_break(document)
    add_page_header(document, "WHO 점검", "이 강의가 누구를 위한 강의인지 구체적으로 정리합니다.")
    add_bullets(document, [
        "이 강의는 정확히 누구를 위한 강의입니까?",
        "이 사람은 지금 어떤 상태에 있습니까?",
        "이 사람이 실제로 겪는 문제는 무엇입니까?",
        "이 강의가 필요한 상황은 무엇입니까?",
    ])
    add_two_column_boxes(document, [
        "대상자의 역할 또는 직무",
        "대상자의 현재 상태",
        "대상자의 실제 문제",
        "이 강의가 필요한 이유",
    ], height_cm=2.8)

    add_page_break(document)
    add_page_header(document, "핵심 콘텐츠 명확화", "내가 실제로 전달하고 싶은 핵심을 압축합니다.")
    add_bullets(document, [
        "핵심 콘텐츠를 1개, 3개, 5개 버전으로 줄이면 무엇입니까?",
        "학습자가 반드시 가져가야 하는 핵심 메시지는 무엇입니까?",
    ])
    add_write_box(document, "이 강의의 핵심 콘텐츠 한 줄", 2.8)
    add_write_box(document, "학습자가 가져가야 하는 핵심 메시지 3개", 4.0)
    add_write_box(document, "덜어낼 수 있는 내용", 3.0)

    add_page_break(document)
    add_page_header(document, "강의 정의 문장", "나는 누구에게 어떠한 내용을 통해 어떠한 변화를 줄 수 있는 강의를 하는가를 한 문장으로 정리합니다.")
    add_write_box(document, "초안", 3.2)
    add_write_box(document, "수정 메모", 2.6)
    add_write_box(document, "최종 문장", 3.2)

    add_page_break(document)
    add_page_header(document, "전달력 플로우", "강의를 Why, What, How, If 흐름으로 재구성합니다.")
    add_illustration_image(document, IMG_FLOW, "전달력 플로우 Why · What · How · If")
    add_two_column_boxes(document, ["Why", "What", "How", "If"], height_cm=2.6)

    add_page_break(document)
    add_page_header(document, "전달력 플로우 기반 1차 구조안", "핵심 문장을 기준으로 1차 구조를 만듭니다.")
    add_two_column_boxes(document, ["도입", "전개", "적용", "정리"], height_cm=3.0)

    add_page_break(document)
    add_page_header(document, "Day 1 피드백 및 Day 2 복기", "첫날 피드백을 정리하고 둘째 날 우선 보완 지점을 잡습니다.")
    add_write_box(document, "Day 1 피드백 정리", 3.5)
    add_write_box(document, "Day 2 우선 보완 포인트", 3.5)
    add_memo_pages(document, 2)

    add_page_break(document)
    add_section_divider(
        document,
        "PART 2",
        "촉진 및 전달",
        "질문, 활동, 에너지 흐름, 전달기술을 다루고 강의를 풍요롭게 하는 기법을 실제 운영 관점에서 정리합니다.",
    )

    add_page_break(document)
    add_page_header(document, "질문 · 활동 설계안", "학습자를 움직이는 질문과 활동을 설계합니다.")
    add_write_box(document, "꼭 넣고 싶은 질문 3개", 3.8)
    add_write_box(document, "꼭 넣고 싶은 활동 2개", 3.8)
    add_write_box(document, "질문과 활동이 연결되는 구간 메모", 3.2)

    add_page_break(document)
    add_page_header(document, "참여도 흐름", "강의 초반에서 후반으로 갈수록 참여 수준이 어떻게 높아지는지 점검합니다.")
    add_illustration_image(document, IMG_ENGAGEMENT, "초반의 낮은 참여 요구에서 후반의 높은 참여 요구로 이어지는 흐름")
    add_write_box(document, "현재 참여 흐름", 3.0)
    add_write_box(document, "수정이 필요한 참여 구간", 3.0)

    add_page_break(document)
    add_page_header(document, "에너지 흐름과 액티비티 설계", "초반과 후반의 에너지 차이에 맞춰 액티비티를 설계합니다.")
    add_illustration_image(document, IMG_ENERGY, "초반에는 강사 에너지 중심, 후반에는 참석자 에너지 중심으로 이동")
    add_blank_write_box(document, 2.8)
    add_blank_write_box(document, 2.8)
    add_blank_write_box(document, 2.8)

    add_page_break(document)
    add_subsection_banner(
        document,
        "강의를 풍요롭게 하는 기법",
        "에너지, 짝토의, 공간 사용, 비언어를 통해 학습이 더 잘 일어나게 만드는 장치를 점검합니다.",
    )

    add_page_break(document)
    add_page_header(document, "도입 몰입을 위한 기술", "강의 초반 몰입을 높이는 장치를 현재 강의 흐름에 맞춰 설계합니다.")
    add_bullets(document, [
        "Story Telling",
        "온도 및 컨디션 체크",
        "Paired Share",
        "주제 관련 Insight Talk",
        "성찰 활동 및 Connect Story",
        "강사 소개 Story Telling",
        "Ground Rule 설정",
    ])
    add_write_box(document, "내 강의 도입에 넣을 장치", 3.0)
    add_write_box(document, "도입 순서 설계", 3.0)

    add_page_break(document)
    add_page_header(document, "Paired Share", "두 사람이 짝을 이루어 짧게 나누는 촉진 방식을 점검합니다.")
    add_illustration_image(document, IMG_PAIRED, "두 사람이 짝을 이루어 짧게 나누는 Paired Share")
    add_write_box(document, "내 강의에 넣을 Paired Share 질문", 2.8)
    add_write_box(document, "진행 방식 메모", 2.8)

    add_page_break(document)
    add_page_header(document, "Paired Share 진행 스크립트", "진행 문구를 미리 써두면 현장에서 훨씬 안정적으로 운영할 수 있습니다.")
    add_bullets(document, [
        "지금부터 둘씩 짝을 지어 이야기해 보겠습니다.",
        "우리가 함께 이야기할 주제는 ______ 입니다.",
        "예를 들면 ______ 이야기를 하면 됩니다.",
        "이야기할 시간은 ○분 드리겠습니다. 순서는 ○가 먼저 시작하면 됩니다.",
        "30초 후 마무리하겠습니다. 혹시 시간이 더 필요하신 분 계신가요?",
        "마무리 후에는 감사와 칭찬으로 연결합니다.",
    ])
    add_write_box(document, "내 수업용 스크립트 초안", 4.0)

    add_page_break(document)
    add_page_header(document, "Gradient · Modeling", "쉬운 것에서 어려운 것, 일반적인 것에서 전문적인 것으로 이동하도록 실습을 설계합니다.")
    add_bullets(document, [
        "실습은 쉬운 것에서 어려운 것으로 이동합니다.",
        "대중적이고 일반적인 것에서 더 전문적인 내용으로 이동합니다.",
        "강사 중심의 디자인에서 학습자 중심의 디자인으로 이동합니다.",
        "사례와 예시는 먼저 구체적으로 제시하고, 긍정/부정 사례를 함께 다루면 학습 밀도가 올라갑니다.",
    ])
    add_two_col_notes(document, "점진적으로 높일 구간", "예시와 사례를 넣을 구간", 3.8)

    add_page_break(document)
    add_page_header(document, "전달기술 점검표와 개인 음성 진단", "전달 방식과 음성 사용 습관을 함께 점검합니다.")
    add_check_table(
        document,
        ["설명", "전환", "강조", "마무리"],
        ["잘 되고 있다", "점검 필요", "수정 필요"],
    )
    add_write_box(document, "전달이 약한 구간", 2.6)
    add_write_box(document, "개인 음성 진단 메모", 2.6)
    add_two_column_boxes(document, ["속도", "강세", "톤", "발음"], height_cm=1.7)

    add_page_break(document)
    add_page_header(document, "동선", "설명과 이동의 관계를 점검합니다.")
    add_illustration_image(document, IMG_MOVEMENT, "스크린 설명, 앞으로 이동, 청중 가까이 이동의 흐름")
    add_write_box(document, "현재 동선 메모", 2.8)
    add_write_box(document, "수정할 동선 설계", 2.8)

    add_page_break(document)
    add_page_header(document, "로케이션 앵커링", "아이디어를 공간 위치와 연결하는 방식을 점검합니다.")
    add_illustration_image(document, IMG_ANCHOR, "큰 무대 위 세 위치를 아이디어와 연결하는 로케이션 앵커링")
    add_write_box(document, "고정 위치로 쓸 구간", 2.8)
    add_write_box(document, "위치와 아이디어 연결 메모", 2.8)

    add_page_break(document)
    add_page_header(document, "비언어 전달", "몸짓과 자세가 어떤 메시지를 주는지 점검합니다.")
    add_illustration_image(document, IMG_NONVERBAL, "Leveler · Placator · Blamer · Thinker · Beggar 제스처")
    add_write_box(document, "내가 자주 쓰는 비언어 패턴", 2.8)
    add_write_box(document, "수정할 비언어 습관", 2.8)
    add_memo_pages(document, 2)

    add_page_break(document)
    add_section_divider(
        document,
        "PART 3",
        "코칭 피드백",
        "미니 시연과 피드백, 수정 반영, 최종 개선안 정리까지 실제 운영 가능한 상태로 마무리합니다.",
    )

    add_page_break(document)
    add_page_header(document, "미니 시연 준비", "어떤 구간을 시연하고 무엇을 점검받을지 먼저 정리합니다.")
    add_write_box(document, "시연 구간", 3.0)
    add_write_box(document, "점검받고 싶은 포인트", 3.0)

    add_page_break(document)
    add_page_header(document, "미니 시연 피드백 기록", "시연 후 받은 피드백을 수정 가능한 언어로 기록합니다.")
    add_two_column_boxes(document, [
        "잘한 점",
        "더 잘할 수 있는 점",
        "메시지 피드백",
        "구조 피드백",
        "촉진 피드백",
        "전달 피드백",
    ], height_cm=2.2)
    add_write_box(document, "즉시 수정 포인트", 2.8)

    add_page_break(document)
    add_page_header(document, "강의 피드백 정리", "피드백을 학습자, 동료, 진행자 관점에서 다시 정리합니다.")
    add_two_col_notes(document, "학습자 관점에서 좋았던 점", "더 잘할 수 있는 점", 3.6)
    add_write_box(document, "다음 운영 전 수정할 점", 3.0)

    add_page_break(document)
    add_page_header(document, "나만의 강의 개성 살리기", "이번 과정에서 드러난 강점과 반응 포인트를 자기 강의의 개성으로 연결합니다.")
    add_bullets(document, [
        "피드백 수집과 반영",
        "이번 과정에서 가장 자신 있었던 순간은 언제였습니까?",
        "학습자나 동료가 긍정적으로 반응한 지점은 무엇입니까?",
        "내가 가장 잘 아는 내용과 가장 즐겁게 전달하는 방식은 어디서 만납니까?",
        "반복 가능한 나만의 강의 패턴으로 만들 수 있는 것은 무엇입니까?",
    ])
    add_write_box(document, "이번 과정에서 드러난 내 강점", 2.8)
    add_write_box(document, "반응이 좋았던 포인트", 2.8)
    add_write_box(document, "반복 가능한 나만의 패턴", 2.8)

    add_page_break(document)
    add_page_header(document, "수정 반영 메모", "피드백 이후 실제로 무엇을 바꿀지 정리합니다.")
    add_write_box(document, "바로 반영할 수정", 3.2)
    add_write_box(document, "다음 운영 전 수정", 3.2)

    add_page_break(document)
    add_page_header(document, "주력 강의 개선안", "과정 전체를 반영해 최종 개선안을 정리합니다.")
    add_two_column_boxes(document, [
        "강의명",
        "대상",
        "강의 정의 문장",
        "핵심 메시지",
        "전달력 플로우 구조 요약",
        "질문/활동 핵심 포인트",
        "전달 수정 포인트",
    ], height_cm=2.2)

    add_page_break(document)
    add_page_header(document, "수정 우선순위와 과정 마무리", "당장 고칠 것과 이후 보완할 것을 구분하고, 과정의 핵심 배움을 정리합니다.")
    add_write_box(document, "바로 수정할 것 3개", 2.6)
    add_write_box(document, "다음 운영 전에 보완할 것 3개", 2.6)
    add_write_box(document, "이번 과정에서 가장 크게 바뀐 기준", 2.6)
    add_write_box(document, "과정 후 2주 안에 실행할 행동", 2.6)
    add_memo_pages(document, 2)

    return document


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    document = build_document()
    document.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    main()
