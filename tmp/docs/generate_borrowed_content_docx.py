from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT_DIR = Path(r"C:\코딩\공개과정\output\doc")
OUT_PATH = OUT_DIR / "주력강의재설계과정_차용콘텐츠_시안.docx"
ILLUS_DIR = Path(r"C:\코딩\공개과정\output\doc\illustrations")
IMG_ENERGY = ILLUS_DIR / "에너지흐름과액티비티설계_삽화_v2.png"
IMG_PAIRED = ILLUS_DIR / "PairedShare_삽화.png"
IMG_MOVEMENT = ILLUS_DIR / "동선_삽화_v2.png"
IMG_ANCHOR = ILLUS_DIR / "로케이션앵커링_삽화_v2.png"
IMG_NONVERBAL = ILLUS_DIR / "비언어전달_삽화_v3.png"

NAVY = RGBColor(23, 55, 117)
GRAY = RGBColor(90, 104, 128)
LIGHT_BLUE = "EAF0FB"
PALE = "FBFCFE"
RULE = RGBColor(210, 220, 236)
FONT_BODY = "Pretendard Variable"
FONT_HEAD = "Pretendard Variable Black"


def apply_font(run, font_name: str) -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:ascii"), font_name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font_name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run._element.rPr.rFonts.set(qn("w:cs"), font_name)


def styled_run(paragraph, text: str, *, size=10.5, bold=False, color=None, font_name=None):
    run = paragraph.add_run(text)
    apply_font(run, font_name or FONT_BODY)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    return run


def set_doc_defaults(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    normal = document.styles["Normal"]
    normal.font.name = FONT_BODY
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT_BODY)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_BODY)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
    normal._element.rPr.rFonts.set(qn("w:cs"), FONT_BODY)
    normal.font.size = Pt(10.5)

    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        style = document.styles[style_name]
        style.font.name = FONT_HEAD
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT_HEAD)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_HEAD)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_HEAD)
        style._element.rPr.rFonts.set(qn("w:cs"), FONT_HEAD)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    styled_run(footer, "주력 강의 재설계 과정 차용 콘텐츠 시안", size=9, color=GRAY)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_borders(cell, *, color: str = "C8D4E8", size: str = "8") -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_row_height(row, height_twips: int) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tr_height = OxmlElement("w:trHeight")
    tr_height.set(qn("w:val"), str(height_twips))
    tr_height.set(qn("w:hRule"), "atLeast")
    tr_pr.append(tr_height)


def add_page_break(document: Document) -> None:
    document.add_page_break()


def add_rule(document: Document) -> None:
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run("_" * 60)
    apply_font(run, FONT_BODY)
    run.font.color.rgb = RULE
    run.font.size = Pt(8)


def add_page_header(document: Document, title: str, subtitle: str | None = None) -> None:
    p = document.add_paragraph()
    styled_run(p, title, size=18, bold=True, color=NAVY, font_name=FONT_HEAD)
    if subtitle:
        sp = document.add_paragraph()
        sp.paragraph_format.space_after = Pt(8)
        styled_run(sp, subtitle, size=10.5, color=GRAY)
    add_rule(document)


def add_section_divider(document: Document, title: str, body: str) -> None:
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    row = table.rows[0]
    set_row_height(row, 1800)
    cell = row.cells[0]
    set_cell_shading(cell, "F2F6FD")
    set_cell_borders(cell)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(30)
    styled_run(p, title + "\n", size=20, bold=True, color=NAVY, font_name=FONT_HEAD)
    styled_run(p, body, size=11, color=GRAY)


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        p = document.add_paragraph(style="List Bullet")
        styled_run(p, item)


def add_write_box(document: Document, label: str, height_cm: float = 3.0) -> None:
    p = document.add_paragraph()
    styled_run(p, label, size=10.5, bold=True, color=NAVY, font_name=FONT_HEAD)
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    row = table.rows[0]
    set_row_height(row, int(height_cm * 567))
    cell = row.cells[0]
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    set_cell_shading(cell, PALE)
    set_cell_borders(cell)
    styled_run(cell.paragraphs[0], " ", size=10)


def add_two_col_notes(document: Document, left_label: str, right_label: str, height_cm: float = 3.0) -> None:
    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    row = table.rows[0]
    set_row_height(row, int(height_cm * 567))
    labels = [left_label, right_label]
    for idx, cell in enumerate(row.cells):
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        set_cell_shading(cell, PALE)
        set_cell_borders(cell)
        p = cell.paragraphs[0]
        styled_run(p, labels[idx] + "\n", size=10, bold=True, color=NAVY, font_name=FONT_HEAD)
        styled_run(p, " ", size=10)


def add_check_table(document: Document, rows: list[str], cols: list[str]) -> None:
    table = document.add_table(rows=1 + len(rows), cols=1 + len(cols))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    set_cell_shading(hdr[0], LIGHT_BLUE)
    set_cell_borders(hdr[0])
    styled_run(hdr[0].paragraphs[0], "점검 항목", size=9.5, bold=True, color=NAVY, font_name=FONT_HEAD)
    for i, col in enumerate(cols, start=1):
        set_cell_shading(hdr[i], LIGHT_BLUE)
        set_cell_borders(hdr[i])
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        styled_run(p, col, size=9.5, bold=True, color=NAVY, font_name=FONT_HEAD)
    for r_idx, row_label in enumerate(rows, start=1):
        row = table.rows[r_idx].cells
        set_cell_shading(row[0], "F6F8FC")
        set_cell_borders(row[0])
        styled_run(row[0].paragraphs[0], row_label, size=9.5)
        for c_idx in range(1, len(cols) + 1):
            set_cell_borders(row[c_idx])
            row[c_idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            styled_run(row[c_idx].paragraphs[0], "□", size=12, color=GRAY)


def add_image(document: Document, image_path: Path, caption: str, width_cm: float = 15.8) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))
    cp = document.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    styled_run(cp, caption, size=9.5, color=GRAY)


def build_document() -> Document:
    d = Document()
    set_doc_defaults(d)

    cover = d.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover.paragraph_format.space_before = Pt(100)
    styled_run(cover, "주력 강의 재설계 과정\n", size=23, bold=True, color=NAVY, font_name=FONT_HEAD)
    styled_run(cover, "차용 콘텐츠 시안", size=16, bold=True, color=NAVY, font_name=FONT_HEAD)

    sub = d.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_before = Pt(18)
    styled_run(sub, "기존 교수법 교안에서 바로 가져와 반영할 수 있는 파트만 별도 정리", size=11, color=GRAY)

    add_page_break(d)
    add_section_divider(d, "Opening과 학습 진입", "기존 교안의 초반 프레임을 현재 과정의 톤에 맞게 차용한 파트입니다.")

    add_page_break(d)
    add_page_header(d, "Opening Frames", "기존 교안 슬라이드 2에서 차용한 도입 프레임")
    add_bullets(d, [
        "새로운 것을 배우고 접할 때 불편하거나 혼란스러운 느낌이 드는 것은 일반적입니다. 이를 변화의 자연스러운 과정으로 받아들일 때 학습이 시작됩니다.",
        "학습은 새로 배운 개념을 실제 적용과 연결할 수 있을 때 가장 빠르게 일어납니다.",
        "강사는 모든 현장 상황의 전문가가 아니므로, 참가자들의 사례와 맥락이 함께 들어와야 학습 효과가 커집니다.",
        "과정 중 참여 수준과 실제 현장 적용 수준은 비례합니다.",
        "소속감을 가지고 성실하게 참여할수록 더 좋은 학습 결과를 얻을 수 있습니다.",
    ])
    add_two_col_notes(d, "현재 과정에 그대로 쓰고 싶은 문구", "전종목 과정 톤에 맞게 바꿔야 할 문구", 3.6)

    add_page_break(d)
    add_page_header(d, "기대사항", "기존 교안 슬라이드 3에서 차용한 질문 구조")
    add_write_box(d, "이번 과정을 통해 가장 얻고 싶은 것은 무엇인가", 3.0)
    add_write_box(d, "이 과정을 통해 도움을 받게 될 사람들은 내가 무엇을 가져오길 기대할까", 3.0)

    add_page_break(d)
    add_section_divider(d, "촉진 및 몰입", "기존 교안의 촉진 기법 파트 중 현재 과정에 직접 연결되는 내용입니다.")

    add_page_break(d)
    add_page_header(d, "도입 몰입을 위한 기술", "기존 교안 슬라이드 12에서 차용한 요소")
    add_bullets(d, [
        "Story Telling",
        "온도 및 컨디션 체크",
        "Paired Share",
        "주제 관련 Insight Talk",
        "성찰 활동 및 Connect Story",
        "강사 소개 Story Telling",
        "Ground Rule 설정 및 보상 시스템 설계",
    ])
    add_two_col_notes(d, "내 강의 도입에 넣고 싶은 장치", "바로 적용 가능한 순서", 3.5)

    add_page_break(d)
    add_page_header(d, "Paired Share 진행 문구", "기존 교안 슬라이드 13의 진행 문장을 작업형으로 전환")
    add_bullets(d, [
        "지금부터 둘씩 짝을 지어 이야기해 보겠습니다.",
        "우리가 함께 이야기할 주제는 ______ 입니다.",
        "예를 들면 ______ 이야기를 하면 됩니다.",
        "이야기할 시간은 ○분 드리겠습니다. 순서는 ○가 먼저 시작하면 됩니다.",
        "30초 후 마무리하겠습니다. 혹시 시간이 더 필요하신 분 계신가요?",
        "마무리 후에는 감사와 칭찬으로 연결합니다.",
    ])
    add_image(d, IMG_PAIRED, "Paired Share")
    add_write_box(d, "내 강의에 맞는 Paired Share 스크립트", 4.0)

    add_page_break(d)
    add_page_header(d, "에너지 흐름과 실습 만들기", "기존 교안 슬라이드 14에서 차용한 핵심 구조")
    add_bullets(d, [
        "강의 실습은 쉬운 것에서 어려운 것으로 이동합니다.",
        "대중적이고 일반적인 것에서 더 전문적인 내용으로 이동합니다.",
        "강사 중심의 디자인에서 학습자 중심의 디자인으로 이동합니다.",
        "사례와 예시는 먼저 구체적으로 제시하고, 긍정/부정 사례를 함께 다루면 학습 밀도가 올라갑니다.",
    ])
    add_image(d, IMG_ENERGY, "에너지 흐름과 액티비티 설계")
    add_two_col_notes(d, "초반에 넣을 에너지 장치", "후반에 유지할 참여 장치", 3.4)

    add_page_break(d)
    add_section_divider(d, "전달 스킬", "공간 사용과 비언어를 현재 과정 언어로 옮긴 파트입니다.")

    add_page_break(d)
    add_page_header(d, "강의 동선", "기존 교안 슬라이드 10에서 차용")
    add_bullets(d, [
        "모든 강의는 동선이 중요합니다.",
        "의미에 따라 동선을 활용합니다.",
        "설명 위치와 예시 위치, 청중과 연결되는 위치를 구분하면 전달 흐름이 더 명확해집니다.",
    ])
    add_image(d, IMG_MOVEMENT, "기본 동선 흐름")
    add_write_box(d, "현재 강의에서 바꿔야 할 동선", 3.3)

    add_page_break(d)
    add_page_header(d, "로케이션 앵커링", "기존 교안 슬라이드 10의 Anchoring 개념을 현재 용어로 정리")
    add_bullets(d, [
        "핵심 아이디어를 서로 다른 위치와 연결하면 청중이 내용의 구획을 더 쉽게 따라옵니다.",
        "질문이 나오면 원래 설명 위치로 돌아가 답하면 흐름이 더 자연스러워집니다.",
    ])
    add_image(d, IMG_ANCHOR, "큰 무대 위 세 위치를 아이디어와 연결하는 로케이션 앵커링")
    add_write_box(d, "내 강의의 핵심 아이디어와 위치 연결", 3.3)

    add_page_break(d)
    add_page_header(d, "비언어적 메시지", "기존 교안 슬라이드 11에서 차용")
    add_bullets(d, [
        "권위, 청유, 지적, 도움의 느낌은 비언어로 먼저 전달되기도 합니다.",
        "말의 메시지와 몸짓의 메시지가 어긋나지 않는지 점검해야 합니다.",
    ])
    add_image(d, IMG_NONVERBAL, "Leveler · Placator · Blamer · Thinker · Beggar")
    add_write_box(d, "내가 자주 쓰는 비언어와 수정 포인트", 3.5)

    add_page_break(d)
    add_section_divider(d, "피드백과 코칭", "기존 교안의 피드백 형식과 개성 탐색 파트를 가져온 영역입니다.")

    add_page_break(d)
    add_page_header(d, "강의 피드백", "기존 교안 슬라이드 15와 41-44의 구조 차용")
    add_two_col_notes(d, "잘한 점", "더 잘할 수 있는 점", 4.0)
    add_check_table(d, ["메시지", "구조", "촉진", "전달"], ["좋음", "보완", "즉시 수정"])

    add_page_break(d)
    add_page_header(d, "나만의 강의 개성 살리기", "기존 교안 슬라이드 20에서 차용")
    add_bullets(d, [
        "피드백 수집과 반영",
        "자기 경험 회고",
        "타인의 시각 비교",
        "반응 포인트 분석",
        "다양한 상황 실험",
        "전문성과 개성의 교차점 찾기",
        "기록과 패턴화",
    ])
    add_write_box(d, "내 강의 개성으로 발전시킬 강점", 3.5)
    add_write_box(d, "반복 가능한 나만의 패턴", 3.5)

    return d


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = build_document()
    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    main()
